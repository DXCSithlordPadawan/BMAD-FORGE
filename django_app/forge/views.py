"""
BMAD Forge Views
Views for document generation and management
"""
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.utils import timezone
from django.db.models import Q, Count

from .models import DocumentTemplate, GeneratedDocument, DocumentActivity
from .forms import DocumentTemplateForm, GenerateDocumentForm
from .document_generator import DocumentGenerator

import json
import logging

logger = logging.getLogger(__name__)


def index(request):
    """
    Homepage view
    """
    recent_documents = GeneratedDocument.objects.filter(
        status__in=['draft', 'review', 'approved']
    )[:5]

    active_templates = DocumentTemplate.objects.filter(is_active=True)
    template_count = active_templates.count()
    document_count = GeneratedDocument.objects.count()

    # Role stats
    role_map = dict(DocumentTemplate.AGENT_ROLES)
    role_counts = (
        active_templates
        .exclude(agent_role='')
        .values('agent_role')
        .annotate(count=Count('id'))
    )
    role_stats = [
        {'label': role_map.get(r['agent_role'], r['agent_role']),
         'key': r['agent_role'], 'count': r['count']}
        for r in role_counts
    ]

    # Phase stats
    phase_map = dict(DocumentTemplate.WORKFLOW_PHASES)
    phase_counts = (
        active_templates
        .exclude(workflow_phase='')
        .values('workflow_phase')
        .annotate(count=Count('id'))
    )
    phase_stats = [
        {'label': phase_map.get(p['workflow_phase'], p['workflow_phase']),
         'key': p['workflow_phase'], 'count': p['count']}
        for p in phase_counts
    ]

    context = {
        'recent_documents': recent_documents,
        'template_count': template_count,
        'document_count': document_count,
        'role_stats': role_stats,
        'phase_stats': phase_stats,
    }

    return render(request, 'forge/index.html', context)


def template_list(request):
    """
    List all active document templates
    """
    templates = DocumentTemplate.objects.filter(is_active=True)

    # Filter by document type if specified
    doc_type = request.GET.get('type')
    if doc_type:
        templates = templates.filter(document_type=doc_type)

    # Filter by workflow phase (before role filter to reduce dataset)
    phase = request.GET.get('phase')
    if phase:
        templates = templates.filter(workflow_phase=phase)

    # Filter by template type (before role filter to reduce dataset)
    template_type = request.GET.get('template_type')
    if template_type:
        templates = templates.filter(template_type=template_type)

    # Search functionality (before role filter to reduce dataset)
    search = request.GET.get('search')
    if search:
        templates = templates.filter(
            Q(name__icontains=search) | Q(description__icontains=search)
        )

    # Filter by agent role
    # NOTE: SQLite doesn't support __contains on JSONField
    # So we evaluate the queryset first, then filter in Python
    role = request.GET.get('role')
    if role:
        # Evaluate queryset to list before checking agent_roles
        all_templates = list(templates)
        role_template_ids = []
        for template in all_templates:
            # Check if role matches agent_role field OR is in agent_roles list
            if template.agent_role == role or role in (template.agent_roles or []):
                role_template_ids.append(template.id)
        # Filter back to queryset with matching IDs
        templates = DocumentTemplate.objects.filter(id__in=role_template_ids)

    context = {
        'templates': templates,
        'document_types': DocumentTemplate.DOCUMENT_TYPES,
        'agent_roles': DocumentTemplate.AGENT_ROLES,
        'workflow_phases': DocumentTemplate.WORKFLOW_PHASES,
        'template_types': DocumentTemplate.TEMPLATE_TYPES,
        'current_type': doc_type,
        'current_role': role,
        'current_phase': phase,
        'current_template_type': template_type,
        'search_query': search,
    }

    return render(request, 'forge/template_list.html', context)


def template_detail(request, template_id):
    """
    View template details
    """
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    # Get recent documents generated from this template
    recent_docs = template.generated_documents.all()[:10]
    
    context = {
        'template': template,
        'recent_documents': recent_docs,
        'variables': template.get_variables(),
    }
    
    return render(request, 'forge/template_detail.html', context)


@login_required
def generate_document(request, template_id):
    """
    Generate a document from a template using wizard interface
    """
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    if request.method == 'POST':
        form = GenerateDocumentForm(request.POST, template=template)
        
        if form.is_valid():
            try:
                # Get form data
                title = form.cleaned_data['title']
                variables = form.get_variables()
                
                # Generate document
                generator = DocumentGenerator(template)
                content = generator.generate(variables)
                
                # Create document record
                document = GeneratedDocument.objects.create(
                    template=template,
                    title=title,
                    content=content,
                    created_by=request.user,
                    status='draft'
                )
                document.set_variables(variables)
                document.save()
                
                # Log activity
                DocumentActivity.objects.create(
                    document=document,
                    activity_type='created',
                    description=f"Document created from template: {template.name}",
                    user=request.user
                )
                
                messages.success(request, f"Document '{title}' generated successfully!")
                return redirect('forge:document_detail', document_id=document.id)
                
            except Exception as e:
                logger.error(f"Document generation failed: {str(e)}")
                messages.error(request, f"Document generation failed: {str(e)}")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = GenerateDocumentForm(template=template)
    
    context = {
        'template': template,
        'form': form,
        'variables': template.get_variables(),
    }
    
    return render(request, 'forge/generate_document_wizard.html', context)


def document_list(request):
    """
    List generated documents with filtering
    """
    documents = GeneratedDocument.objects.select_related('template', 'created_by')
    
    # Filter by status
    status = request.GET.get('status')
    if status:
        documents = documents.filter(status=status)
    
    # Filter by template
    template_id = request.GET.get('template')
    if template_id:
        documents = documents.filter(template_id=template_id)
    
    # Search
    search = request.GET.get('search')
    if search:
        documents = documents.filter(
            Q(title__icontains=search) | 
            Q(template__name__icontains=search)
        )
    
    context = {
        'documents': documents,
        'status_choices': GeneratedDocument.STATUS_CHOICES,
        'current_status': status,
        'search_query': search,
    }
    
    return render(request, 'forge/document_list.html', context)


def document_detail(request, document_id):
    """
    View document details and preview
    """
    document = get_object_or_404(
        GeneratedDocument.objects.select_related('template', 'created_by', 'reviewed_by'),
        id=document_id
    )
    
    # Get document activity history
    activities = document.activities.select_related('user')[:20]
    
    # Get document versions
    if document.parent_document:
        versions = document.parent_document.versions.all()
    else:
        versions = document.versions.all()
    
    context = {
        'document': document,
        'activities': activities,
        'versions': versions,
        'variables': document.get_variables(),
    }
    
    return render(request, 'forge/document_detail.html', context)


@login_required
@require_http_methods(["POST"])
def document_approve(request, document_id):
    """
    Approve a document
    """
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    try:
        document.approve(request.user)
        
        DocumentActivity.objects.create(
            document=document,
            activity_type='approved',
            description=f"Document approved by {request.user.get_full_name() or request.user.username}",
            user=request.user
        )
        
        messages.success(request, f"Document '{document.title}' approved successfully!")
    except Exception as e:
        logger.error(f"Document approval failed: {str(e)}")
        messages.error(request, f"Approval failed: {str(e)}")
    
    return redirect('forge:document_detail', document_id=document.id)


@login_required
@require_http_methods(["POST", "DELETE"])
def document_delete(request, document_id):
    """
    Delete a document
    """
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    # Store document title for success message
    document_title = document.title
    
    try:
        # Log deletion activity before deleting
        DocumentActivity.objects.create(
            document=document,
            activity_type='archived',
            description=f"Document deleted by {request.user.get_full_name() or request.user.username}",
            user=request.user
        )
        
        # Delete the document
        document.delete()
        
        messages.success(request, f"Document '{document_title}' deleted successfully!")
        return redirect('forge:document_list')
        
    except Exception as e:
        logger.error(f"Document deletion failed: {str(e)}")
        messages.error(request, f"Deletion failed: {str(e)}")
        return redirect('forge:document_detail', document_id=document.id)


@login_required
def document_download(request, document_id, format='html'):
    """
    Download document in specified format
    Supports: html, markdown, docx, pdf
    """
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    # Log download activity
    DocumentActivity.objects.create(
        document=document,
        activity_type='downloaded',
        description=f"Document downloaded in {format.upper()} format",
        user=request.user
    )
    
    if format == 'html':
        response = HttpResponse(document.content, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{document.file_name or document.title}.html"'
        return response
    
    elif format == 'markdown' or format == 'md':
        # Convert HTML to proper Markdown format
        import re
        from html.parser import HTMLParser
        from html import unescape
        
        class HTMLToMarkdownConverter(HTMLParser):
            """Proper HTML to Markdown converter"""
            def __init__(self):
                super().__init__()
                self.markdown_lines = []
                self.current_line = []
                self.in_heading = False
                self.heading_level = 0
                self.in_paragraph = False
                self.in_list = False
                self.in_list_item = False
                self.in_bold = False
                self.in_italic = False
                self.in_code = False
                self.in_pre = False
                self.list_level = 0
                self.pre_content = []
                
            def handle_starttag(self, tag, attrs):
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.flush_line()
                    self.in_heading = True
                    self.heading_level = int(tag[1])
                    self.current_line.append('#' * self.heading_level + ' ')
                    
                elif tag == 'p':
                    self.flush_line()
                    self.in_paragraph = True
                    
                elif tag == 'br':
                    self.current_line.append('  \n')  # Markdown line break
                    
                elif tag in ['ul', 'ol']:
                    self.flush_line()
                    self.in_list = True
                    self.list_level += 1
                    
                elif tag == 'li':
                    self.flush_line()
                    self.in_list_item = True
                    indent = '  ' * (self.list_level - 1)
                    self.current_line.append(f"{indent}- ")
                    
                elif tag in ['strong', 'b']:
                    self.in_bold = True
                    self.current_line.append('**')
                    
                elif tag in ['em', 'i']:
                    self.in_italic = True
                    self.current_line.append('*')
                    
                elif tag == 'code':
                    self.in_code = True
                    if not self.in_pre:
                        self.current_line.append('`')
                        
                elif tag == 'pre':
                    self.flush_line()
                    self.in_pre = True
                    self.pre_content = []
                    self.current_line.append('```\n')
                    
                elif tag == 'a':
                    # Extract href
                    href = ''
                    for attr, value in attrs:
                        if attr == 'href':
                            href = value
                            break
                    self.current_line.append('[')
                    
                elif tag == 'hr':
                    self.flush_line()
                    self.markdown_lines.append('---')
                    
                elif tag == 'blockquote':
                    self.flush_line()
                    self.current_line.append('> ')
                    
            def handle_endtag(self, tag):
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.in_heading = False
                    self.flush_line()
                    self.markdown_lines.append('')  # Blank line after heading
                    
                elif tag == 'p':
                    self.in_paragraph = False
                    self.flush_line()
                    self.markdown_lines.append('')  # Blank line after paragraph
                    
                elif tag in ['ul', 'ol']:
                    self.list_level -= 1
                    if self.list_level == 0:
                        self.in_list = False
                        self.markdown_lines.append('')  # Blank line after list
                        
                elif tag == 'li':
                    self.in_list_item = False
                    self.flush_line()
                    
                elif tag in ['strong', 'b']:
                    self.in_bold = False
                    self.current_line.append('**')
                    
                elif tag in ['em', 'i']:
                    self.in_italic = False
                    self.current_line.append('*')
                    
                elif tag == 'code':
                    self.in_code = False
                    if not self.in_pre:
                        self.current_line.append('`')
                        
                elif tag == 'pre':
                    self.in_pre = False
                    self.current_line.append('```')
                    self.flush_line()
                    self.markdown_lines.append('')
                    
            def handle_data(self, data):
                if data.strip():
                    # Unescape HTML entities
                    clean_data = unescape(data)
                    # Preserve whitespace in pre blocks
                    if self.in_pre:
                        self.current_line.append(clean_data)
                    else:
                        # Normalize whitespace
                        clean_data = ' '.join(clean_data.split())
                        if clean_data:
                            self.current_line.append(clean_data)
                            
            def flush_line(self):
                """Flush current line to output"""
                if self.current_line:
                    line = ''.join(self.current_line).strip()
                    if line:
                        self.markdown_lines.append(line)
                    self.current_line = []
                    
            def get_markdown(self):
                """Get final markdown output"""
                self.flush_line()
                # Join lines and clean up excessive blank lines
                markdown = '\n'.join(self.markdown_lines)
                # Remove more than 2 consecutive newlines
                markdown = re.sub(r'\n{3,}', '\n\n', markdown)
                return markdown.strip() + '\n'
        
        # Convert HTML to Markdown
        try:
            converter = HTMLToMarkdownConverter()
            converter.feed(document.content)
            markdown_content = converter.get_markdown()
            
            # Add document metadata as frontmatter
            frontmatter = f"""---
title: {document.title}
template: {document.template.name}
created: {document.created_at.strftime('%Y-%m-%d')}
author: {document.created_by.get_full_name() or document.created_by.username if document.created_by else 'Unknown'}
status: {document.get_status_display()}
version: {document.version}
---

"""
            markdown_content = frontmatter + markdown_content
            
        except Exception as e:
            logger.error(f"Markdown conversion failed: {str(e)}")
            # Fallback to simple HTML stripping
            markdown_content = re.sub(r'<[^>]+>', '', document.content)
            markdown_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', markdown_content)
        
        response = HttpResponse(markdown_content, content_type='text/markdown')
        filename = f"{document.file_name or document.title}.md"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    elif format == 'docx':
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import re
            from html.parser import HTMLParser
            
            # Create a new Document
            doc = Document()
            
            # Add document title
            title = doc.add_heading(document.title, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Generated: {document.created_at.strftime('%B %d, %Y')}")
            doc.add_paragraph(f"Template: {document.template.name}")
            if document.created_by:
                doc.add_paragraph(f"Created by: {document.created_by.get_full_name() or document.created_by.username}")
            doc.add_paragraph()  # Blank line
            
            # Parse and add content
            # Simple HTML to DOCX conversion
            class HTMLToDocx(HTMLParser):
                def __init__(self, doc):
                    super().__init__()
                    self.doc = doc
                    self.current_paragraph = None
                    self.in_heading = False
                    self.heading_level = 1
                    self.in_bold = False
                    self.in_italic = False
                    self.in_list = False
                    self.list_items = []
                    
                def handle_starttag(self, tag, attrs):
                    if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self.in_heading = True
                        self.heading_level = int(tag[1])
                    elif tag == 'p':
                        self.current_paragraph = self.doc.add_paragraph()
                    elif tag in ['ul', 'ol']:
                        self.in_list = True
                    elif tag == 'li':
                        self.list_items.append([])
                    elif tag == 'strong' or tag == 'b':
                        self.in_bold = True
                    elif tag == 'em' or tag == 'i':
                        self.in_italic = True
                    elif tag == 'br':
                        if self.current_paragraph:
                            self.current_paragraph.add_run('\n')
                            
                def handle_endtag(self, tag):
                    if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self.in_heading = False
                    elif tag == 'p':
                        self.current_paragraph = None
                    elif tag in ['ul', 'ol']:
                        # Add list items
                        for item_text in self.list_items:
                            text = ''.join(item_text)
                            if text.strip():
                                self.doc.add_paragraph(text, style='List Bullet')
                        self.list_items = []
                        self.in_list = False
                    elif tag == 'strong' or tag == 'b':
                        self.in_bold = False
                    elif tag == 'em' or tag == 'i':
                        self.in_italic = False
                        
                def handle_data(self, data):
                    if not data.strip():
                        return
                        
                    if self.in_heading:
                        self.doc.add_heading(data.strip(), level=self.heading_level)
                    elif self.in_list and self.list_items:
                        self.list_items[-1].append(data)
                    elif self.current_paragraph:
                        run = self.current_paragraph.add_run(data)
                        if self.in_bold:
                            run.bold = True
                        if self.in_italic:
                            run.italic = True
                    else:
                        # No current paragraph, create one
                        para = self.doc.add_paragraph()
                        run = para.add_run(data)
                        if self.in_bold:
                            run.bold = True
                        if self.in_italic:
                            run.italic = True
            
            # Parse HTML content
            parser = HTMLToDocx(doc)
            try:
                parser.feed(document.content)
            except Exception as e:
                # If HTML parsing fails, add content as plain text
                logger.warning(f"HTML parsing failed, using plain text: {e}")
                # Strip HTML tags and add as plain text
                plain_text = re.sub(r'<[^>]+>', '', document.content)
                for paragraph in plain_text.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Save to BytesIO
            from io import BytesIO
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            # Create response
            response = HttpResponse(
                doc_io.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"{document.file_name or document.title}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except ImportError as e:
            logger.error(f"python-docx not installed: {e}")
            messages.error(request, "DOCX export requires python-docx package. Please install it.")
            return redirect('forge:document_detail', document_id=document.id)
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            messages.error(request, f"DOCX generation failed: {str(e)}")
            return redirect('forge:document_detail', document_id=document.id)
    
    elif format == 'pdf':
        # PDF generation would require additional library (e.g., weasyprint, reportlab)
        messages.warning(request, "PDF export not yet implemented. Consider using DOCX export and converting to PDF in Word.")
        return redirect('forge:document_detail', document_id=document.id)
    
    else:
        messages.error(request, f"Unsupported format: {format}")
        return redirect('forge:document_detail', document_id=document.id)


@require_http_methods(["GET"])
def api_template_variables(request, template_id):
    """
    API endpoint to get template variables
    """
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    return JsonResponse({
        'success': True,
        'template_id': template.id,
        'template_name': template.name,
        'variables': template.get_variables()
    })


@require_http_methods(["POST"])
def api_validate_document(request):
    """
    API endpoint to validate document data before generation
    """
    try:
        data = json.loads(request.body)
        template_id = data.get('template_id')
        variables = data.get('variables', {})
        
        template = get_object_or_404(DocumentTemplate, id=template_id)
        required_vars = template.get_variables()
        
        # Validate required variables
        missing = []
        for var_name, var_config in required_vars.items():
            if var_config.get('required', False) and not variables.get(var_name):
                missing.append(var_name)
        
        if missing:
            return JsonResponse({
                'success': False,
                'errors': {
                    'missing_variables': missing
                }
            }, status=400)
        
        return JsonResponse({
            'success': True,
            'message': 'Validation passed'
        })
        
    except Exception as e:
        logger.error(f"Validation error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)
